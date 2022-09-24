from functools import partial
# import re
# import ctypes
import sys
from warnings import catch_warnings
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import QLabel,QLineEdit,QPushButton,QMessageBox,QFileDialog,QApplication,QWidget
from binascii import a2b_base64
# from re import X
# import webbrowser
from docx import Document
import xlwt
import logging
import os


#界面样式和调用模块
class MyWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()
        self.setFixedSize(self.width(), self.height())

    def init_ui(self):
        self.setWindowTitle("Word转Excel")
        # 更改当前Widge的宽高
        self.resize(400, 100)
        

        # 下面创建一个Label，然后调用方法指定父类
        label1 = QLabel("Word 路径: ", self)
        label1.setGeometry(20, 20, 80, 30)  # 显示位置与大小 ： x, y , w, h

        # label2 = QLabel("日期名字: ", self)
        # label2.setGeometry(20, 46, 80, 30) 

        # 文本框
        self.edit = QLineEdit(self)
        self.edit.setPlaceholderText("请输入路径")
        self.edit.setGeometry(80, 25, 250, 20)
        # edit.setText("aaaa")

        # 文本框
        # self.edit1 = QLineEdit(self)
        # self.edit1.setPlaceholderText("请输入名字前两个字")
        # self.edit1.setGeometry(80, 50, 150, 20)

        # 创建一个按钮 执行按钮
        btn = QPushButton("执行", self)
        btn.setGeometry(150, 50, 100, 30)  # 设置窗口位置、宽高
        # 将按钮被点击时触发的信号与我们定义的函数（方法）进行绑定
        # 注意：这里没有()，即写函数的名字，而不是名字()
        # btn.clicked.connect(self.DataAnalysis)
        btn.clicked.connect(partial(self.click_my_btn))

        #选择文件按钮
        btn1 = QPushButton("选择文件", self)
        btn1.setGeometry(330, 25, 60, 21)  # 设置窗口位置、宽高
        btn1.clicked.connect(partial(self.msg))

        #关于按钮
        btn_about = QPushButton("About", self)
        btn_about.setGeometry(345, 75, 50, 20)  # 设置窗口位置、宽高
        btn_about.clicked.connect(partial(self.msg_about))

        #关于按钮
        bth_help = QPushButton("Help", self)
        bth_help.setGeometry(295, 75, 50, 20)  # 设置窗口位置、宽高
        bth_help.clicked.connect(partial(self.msg_help))

    #触发
    def click_my_btn(self):
        # 槽函数，点击按钮则调用该函数
        # 这里的参数正好是信号发出，传递的参数
        # name=self.edit1.text()
        rou=self.edit.text()
        k=DataAnalysis(rou)
        # print(k)
        if k!="":
            if k=="route_err":
                QMessageBox.critical(self,"出错了","请输入正确路径或名字(っ °Д °;)っ")
            elif k=="excel_err":
                QMessageBox.critical(self,"出错了","请关闭对应excel再进行操作(っ °Д °;)っ")
            elif "add_par" in k:
                if k=="add_par":
                    QMessageBox.warning(self,"警告","文件已生成，已使用行数不够自动补齐功能，但可能会出现未知的错误，请自行检查对应excel文档！如出现错误请参考帮助！")
                else:
                    QMessageBox.warning(self,"警告","文件已生成，已使用行数不够自动补齐功能，但可能会出现未知的错误，请自行检查对应excel文档！如出现错误请参考帮助！"
                                                        "\n"
                                                        "Word文件第"+k.split("r")[1]+"行有参数为空值，如有问题请自行检查o((>ω< ))o")
            else:
                QMessageBox.warning(self,"完成","文件已生成在Word同一路径下"
                                                    "\n"
                                                            "Word文件第"+k+"行有参数为空值，如有问题请自行检查o((>ω< ))o")
        else:
            QMessageBox.information(self,"完成","文件已生成在Word同一路径下(。・ω・。)")
            # print(e)

    def msg(self):
        try:
            directory = QFileDialog.getOpenFileName(self, "选取文件","./", "All Files (*);;Text Files (*.docx)")
            self.edit.setText(directory[0])
        except Exception as e:
            log.error(e)
            # print("错误")

    #关于
    def msg_about(self):
        QMessageBox.about(self, "About",
                       "版本：v0.4"
                       "\n"
                       "作者：https://github.com/LuckyCatL"
                       "\n"
                       "最后更新日期：22-09-22 12:30"
                       "\n"
                       "简介：新手学习python随便写的，如有错误还请包涵"
                       "\n"
                       "         本应用只适用于特定的Word文档"
                       "\n"
                       "更新：1.添加了行数不够自动补齐"
                       "\n"
                       "         2.对不同空格的识别"
                       "\n"
                       "         3.添加了对错误的捕获，对空参数输出“-”并提醒"
                       "\n"
                       "         4.增加了日志模块，对错误进行捕获并保存在./log/error.log文件中"
                       )
    def msg_help(self):
        QMessageBox.information(self, "Help",   "如果使用自动补齐功能出现了错误"
                                                "\n"
                                                "解决方案：请将word文档的“手动换行符”替换为“段落标记符”"      
                                                )

#数据处理存储模块
def DataAnalysis(route):
        # print(name)
        errlist=""  #错误捕获参数
        #分析传进来的路径参数route
        route1=route.rsplit("/",1)
        route_name=route1[1].split(".")[0]
        # print(route_name)

        book = xlwt.Workbook(encoding='utf-8',style_compression=0)  #创建excel
        sheet = book.add_sheet('sheet1',cell_overwrite_ok=True)     #创建名为。。。的sheet表
        col = ('时间','姓名','性别','民族','出生日期','住址','身份证号','电话','去向')     #自定义列名
        first_col=sheet.col(5)
        first_col.width=256*40
        second_col=sheet.col(6)
        second_col.width=256*20
        third_col=sheet.col(7)
        third_col.width=256*20     #定义列宽

        for i in range(0,9): 
            sheet.write(0,i,col[i])

        try:
            document = Document(route) #打开文档
        except Exception as e:
            errlist="route_err"
            log.error(e)
            return errlist
        
        # print(document.paragraphs[0].text)
        #自动补齐功能
        if len(document.paragraphs)%10!=0:
            for num in range(10-(len(document.paragraphs)%10)):
                document.add_paragraph(text="", style=None)
                errlist="add_par"
        #获取所有段落
        # print(int(len(document.paragraphs)/11))
        # print(len(document.paragraphs))   #段落数
        x=0
        
        list = [["a" for j in range(0, 9)] for i in range(int(len(document.paragraphs)/10))]    #初始化二维数组
        # print(list)

        for i in document.paragraphs:
            x=x+1
            a=x//10 #计数
            # print(a)
            # print(i.text)
            try:
                if i.text[0:2]==document.paragraphs[0].text[0:2]:
                    # print(i.text.split("\u00A0")[1])
                    if "\u00A0" in i.text:
                        if i.text.split("\u00A0",1)[1]=="":
                            errlist=errlist+str(x)+","
                            list[a][x%10-1]="-"
                        else:
                            list[a][0]=i.text.split("\u00A0",1)[1]
                    elif " " in i.text:
                        if i.text.split(" ",1)[1]=="":
                            errlist=errlist+str(x)+","
                            list[a][x%10-1]="-"
                        else:
                            list[a][0]=i.text.split(" ",1)[1]
                    # print(list[a][0])
                elif i.text[0:2]=="姓名":
                    # print(i.text.split(':')[0])
                    if i.text.split(':')[1]=="":
                        list[a][1]="-"
                        errlist=errlist+str(x)+","
                    else:
                        list[a][1]=i.text.split(':')[1]
                elif i.text[0:2]=="性别":
                    # print(i.text.split(':')[1])
                    if i.text.split(':')[1]=="":
                        list[a][2]="-"
                        errlist=errlist+str(x)+","
                    else:
                        list[a][2]=i.text.split(':')[1]
                elif i.text[0:2]=="民族":
                    # print(i.text.split(':')[1])
                    if i.text.split(':')[1]=="":
                        list[a][3]="-"
                        errlist=errlist+str(x)+","
                    else:
                        list[a][3]=i.text.split(':')[1]
                elif i.text[0:4]=="出生日期":
                    # print(i.text.split(':')[1])
                    if i.text.split(':')[1]=="":
                        list[a][4]="-"
                        errlist=errlist+str(x)+","
                    else:
                        list[a][4]=i.text.split(':')[1]
                elif i.text[0:2]=="住址":
                    # print(i.text.split(':')[1])
                    if i.text.split(':')[1]=="":
                        list[a][5]="-"
                        errlist=errlist+str(x)+","
                    else:
                        list[a][5]=i.text.split(':')[1]
                elif i.text[0:5]=="身份证号码":
                    # print(i.text.split(':')[1])
                    if i.text.split(':')[1]=="":
                        list[a][6]="-"
                        errlist=errlist+str(x)+","
                    else:
                        list[a][6]=i.text.split(':')[1]
                elif i.text[0:3]=="有效期":
                    # print(i.text.split(' ')[2])
                    if "\u00A0" in i.text:                          #参数中含有“\u00A0”的情况
                        try:
                            phone=i.text.split("\u00A0")[0].split(":")[1]
                            # print(i.text.split("\u00A0")[0])
                        except Exception as e:
                            log.error(e)
                            # print(e)
                            errlist=errlist+str(x)+","
                            phone="-"
                        try:
                            adress=i.text.split("\u00A0")[2]
                            # print(adress)
                        except Exception as e:
                            log.error(e)
                            errlist=errlist+str(x)+","
                            adress="-"
                        c =[phone,adress]
                        # print(c)
                        for n in range(2):
                            list[a][n+7]=c[n]
                            # print(list[a][n+7])
                    elif " " in i.text:                             #参数中含有空格的情况
                        try:
                            phone=i.text.split(" ")[0].split(":")[1]
                        except Exception as e:
                            log.error(e)
                            errlist=errlist+str(x)+","
                            phone="-"
                        try:
                            adress=i.text.split(" ")[2]
                        except Exception as e:
                            log.error(e)
                            errlist=errlist+str(x)+","
                            adress="-"
                        c =[phone,adress]
                        # print(c)
                        for n in range(2):
                            list[a][n+7]=c[n]
                    else:                                           #参数中什么都不含有的情况
                        errlist=errlist+str(x)+","
                        phone=i.text.split(":")[1]
                        adress="-"
                        c =[phone,adress]
                        for n in range(2):
                            list[a][n+7]=c[n]

            except Exception as e:
                log.error(e)
        
        # print(list)
        try:
            for i in range(len(list)):
                for j in range(len(list[i])):
                    sheet.write(i+1, j, list[i][j])     #填入数据

            savepath=route1[0]+"/"+route_name+".xlsx"

            book.save(savepath)     #保存数据
        except Exception as e:
            errlist="excel_err"
            log.error(e)
        return errlist

#日志模块
def init_log(path):
    if os.path.exists(path):
        mode = 'a'
    else:
        mode = 'w'
    logging.basicConfig(        # 针对 basicConfig 进行配置(basicConfig 其实就是对 logging 模块进行动态的调整，之后可以直接使用)
        level=logging.INFO,     # INFO 等级以下的日志不会被记录
        format='[%(asctime)s] %(filename)s[line:%(lineno)d] %(levelname)s %(message)s',    # 日志输出格式
        filename='.\log\error.log',    # 日志存放路径(存放在当前相对路径)
        filemode=mode,          # 输入模式；如果当前我们文件已经存在，可以使用 'a' 模式替代 'w' 模式
                                # 与文件写入的模式相似，'w' 模式为没有文件时创建文件；'a' 模式为追加内容写入日志文件
    )
 
    return logging

current_path = os.getcwd()
path = os.path.join(current_path, '.\log\error.log')
 
log = init_log(path)            # 初始化返回的 init_log() 函数 , 其实就是 return logging

if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon("./server-icon.png"))
    w = MyWindow()
    w.show()

    app.exec()
